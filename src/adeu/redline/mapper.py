from copy import deepcopy
from dataclasses import dataclass
from typing import List, Tuple

import structlog
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = structlog.get_logger(__name__)

@dataclass
class TextSpan:
    start: int
    end: int
    text: str
    run: Run
    paragraph: Paragraph

class DocumentMapper:
    """
    Maps a linear text representation of the document back to the specific
    XML Run elements that contain that text. Handles splitting runs when
    text matches cross run boundaries.
    """
    def __init__(self, doc: DocumentObject):
        self.doc = doc
        self.full_text = ""
        self.spans: List[TextSpan] = []
        self._build_map()

    def _build_map(self):
        current_offset = 0
        self.spans = []
        self.full_text = ""

        # Iterate over all paragraphs (Body + Tables would need recursion in full implementation)
        # For MVP, we stick to body paragraphs + simple table support if needed
        all_paragraphs = list(self.doc.paragraphs)
        
        logger.debug(f"Building map for {len(all_paragraphs)} paragraphs")
        # Add table paragraphs to the map
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for p in all_paragraphs:
            for run in p.runs:
                text_len = len(run.text)
                if text_len == 0:
                    continue

                span = TextSpan(
                    start=current_offset,
                    end=current_offset + text_len,
                    text=run.text,
                    run=run,
                    paragraph=p,
                )
                self.spans.append(span)
                self.full_text += run.text
                current_offset += text_len

            # Add virtual newlines to match ingestion logic (usually \n\n for markdown paragraphs)
            # NOTE: ingest.py joins with "\n\n", so we must match that here.
            self.full_text += "\n\n"
            current_offset += 2

    def find_target_runs(self, target_text: str) -> List[Run]:
        """
        Locates the runs corresponding to target_text. 
        Splits runs if the target starts/ends in the middle of a run.
        Rebuilds the map after splitting.
        """
        logger.debug(f"Searching for target: '{target_text[:50]}...' (Length: {len(target_text)})")
        start_idx = self.full_text.find(target_text)
        if start_idx == -1:
            # Try to handle potential whitespace mismatches roughly?
            # For now, strict match.
            logger.warning("Target text not found", target=target_text[:20])
            return []

        end_idx = start_idx + len(target_text)
        logger.debug(f"Found match at indices [{start_idx}:{end_idx}]")

        affected_spans = [
            s for s in self.spans if s.end > start_idx and s.start < end_idx
        ]

        if not affected_spans:
            return []

        working_runs = [s.run for s in affected_spans]
        
        # We need to track if we modified the DOM, so we can rebuild the map
        dom_modified = False

        # 1. Handle Left Boundary
        first_span = affected_spans[0]
        local_start = start_idx - first_span.start

        if local_start > 0:
            _, right_run = self._split_run_at_index(working_runs[0], local_start)
            working_runs[0] = right_run
            dom_modified = True

        # 2. Handle Right Boundary
        # We re-calculate length because the first split might have changed things effectively
        # But logically, the text content sum hasn't changed, just the run wrappers.
        # Use span indices to handle virtual newlines (\n\n) correctly
        last_span = affected_spans[-1]
        extra_len = last_span.end - end_idx

        if extra_len > 0:
            last_run = working_runs[-1]
            split_point = len(last_run.text) - extra_len
            left_run, _ = self._split_run_at_index(last_run, split_point)
            working_runs[-1] = left_run
            dom_modified = True
            
        if dom_modified:
            self._build_map()

        return working_runs

    def _split_run_at_index(self, run: Run, split_index: int) -> Tuple[Run, Run]:
        """
        Splits a run at a specific character index using raw XML manipulation.
        Returns (left_run_wrapper, right_run_wrapper).
        """
        text = run.text
        left_text = text[:split_index]
        right_text = text[split_index:]

        # 1. Update original run (Left)
        run.text = left_text

        # 2. Create new element for Right
        new_r_element = deepcopy(run._element)

        # Clean existing text nodes
        t_list = new_r_element.findall(qn("w:t"))
        for t in t_list:
            new_r_element.remove(t)

        new_t = OxmlElement("w:t")
        new_t.text = right_text
        if right_text.strip() != right_text:
            new_t.set(qn("xml:space"), "preserve")
        new_r_element.append(new_t)

        # 3. Insert into tree
        parent = run._element.getparent()
        current_idx = parent.index(run._element)
        parent.insert(current_idx + 1, new_r_element)

        # 4. Wrap
        new_run = Run(new_r_element, run._parent)

        return run, new_run