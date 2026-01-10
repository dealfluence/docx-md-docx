import io
import pytest
from docx import Document
from adeu.models import DocumentEdit, EditOperationType
from adeu.redline.engine import RedlineEngine
from adeu.redline.mapper import DocumentMapper
from adeu.ingest import extract_text_from_stream

def test_full_roundtrip_workflow(simple_docx_stream):
    """
    Tests the full lifecycle:
    1. Ingest DOCX -> Text
    2. Simulate LLM creating an Edit
    3. Apply Edit -> Redlined DOCX
    """
    # 1. Ingestion
    extracted_text = extract_text_from_stream(simple_docx_stream)
    assert "Contract Agreement" in extracted_text
    assert "Seller" in extracted_text

    # 2. Simulate LLM Response (ComplianceEdit)
    # Let's change "Seller" to "Vendor"
    edit = DocumentEdit(
        operation=EditOperationType.MODIFICATION,
        target_text="Seller",
        new_text="Vendor",
        comment="Standardizing terminology."
    )

    # 3. Injection (Redlining)
    # Reset stream pointer for the engine (simulating fresh read)
    simple_docx_stream.seek(0)
    engine = RedlineEngine(simple_docx_stream)
    
    engine.apply_edits([edit])
    
    result_stream = engine.save_to_stream()
    
    # 4. Verification
    # Parse the result and check for Track Changes XML
    doc = Document(result_stream)
    
    # Check text content (python-docx .text property usually shows the "current" view, 
    # but exact behavior depends on how it parses w:ins/w:del. 
    # We check the XML directly for the most robust test.)
    
    xml_content = doc.element.xml
    
    # Check for Deletion tag
    assert "w:del" in xml_content
    assert "<w:delText>Seller</w:delText>" in xml_content
    
    # Check for Insertion tag
    assert "w:ins" in xml_content
    assert "<w:t>Vendor</w:t>" in xml_content

def test_split_run_behavior():
    """
    Tests that the engine correctly splits a run when the target text 
    is in the middle of a sentence.
    """
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("The quick brown fox.")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    edit = DocumentEdit(
        operation=EditOperationType.DELETION,
        target_text="brown", # Middle of the run
        new_text=None
    )
    
    engine = RedlineEngine(stream)
    engine.apply_edits([edit])
    
    result_stream = engine.save_to_stream()
    
    # Check XML
    doc = Document(result_stream)
    xml_content = doc.element.xml
    
    # "brown" should be wrapped in delete
    # "The quick " and " fox." should remain as runs (or be split out)
    assert "<w:delText>brown</w:delText>" in xml_content

def test_insertion_spacing_between_complex_runs():
    """
    Reproduction test for 'Extra Space' or misplaced insertion bug.
    Scenario: Text is 'ARTICLE3 FEES' split into runs ['ARTICLE', '3 ', 'FEES'].
    Action: Insert ' ' after 'ARTICLE' and ' ' after '3 '.
    Expected: 'ARTICLE' -> INS(' ') -> '3 ' -> INS(' ') -> 'FEES'
    """
    doc = Document()
    p = doc.add_paragraph()
    
    # Create fragmented runs (simulate bold/not bold to prevent coalescing)
    r1 = p.add_run("ARTICLE")
    r1.bold = True
    r2 = p.add_run("3 ") # Note the space
    r2.bold = False
    r3 = p.add_run("FEES")
    r3.bold = True
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    # Edit 1: Insert space after ARTICLE
    edit1 = DocumentEdit(
        operation=EditOperationType.INSERTION,
        target_text="ARTICLE", 
        new_text=" "
    )
    
    # Edit 2: Insert space after "3 " (which ends in space)
    # This mimics the "3 FEES" edit if the anchor matched "3 FEES"
    edit2 = DocumentEdit(
        operation=EditOperationType.INSERTION,
        target_text="3 ", 
        new_text=" "
    )
    
    engine = RedlineEngine(stream)
    engine.apply_edits([edit1, edit2])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    # 1. Verify "3 " is still there
    assert ">3 </w:t>" in xml
    # 2. Verify "FEES" is there
    assert ">FEES</w:t>" in xml
    
    # Verify Order
    # ARTICLE -> INS(1) -> 3  -> INS(2) -> FEES
    idx_art = xml.find(">ARTICLE</w:t>")
    idx_3 = xml.find(">3 </w:t>")
    idx_fees = xml.find(">FEES</w:t>")
    
    assert idx_art < idx_3, "ARTICLE before 3"
    assert idx_3 < idx_fees, "3 before FEES"
    
    # Check that there is an insertion between ART and 3
    segment_1 = xml[idx_art:idx_3]
    assert "w:ins" in segment_1, "Missing insertion between ARTICLE and 3"
    
    # Check that there is an insertion between 3 and FEES
    segment_2 = xml[idx_3:idx_fees]
    assert "w:ins" in segment_2, "Missing insertion between 3 and FEES"

def test_insertion_splits_coalesced_run():
    """
    Tests inserting text into a run that needs to be split.
    Original: Run("ARTICLE3") (Coalesced)
    Edit: Insert " " after "ARTICLE".
    Expected: Run("ARTICLE"), Ins(" "), Run("3")
    """
    doc = Document()
    p = doc.add_paragraph()
    # Create one run
    run = p.add_run("ARTICLE3")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    edit = DocumentEdit(
        operation=EditOperationType.INSERTION,
        target_text="ARTICLE",
        new_text=" "
    )
    
    engine = RedlineEngine(stream)
    engine.apply_edits([edit])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    # Check order: ARTICLE -> INS -> 3
    idx_art = xml.find(">ARTICLE</w:t>")
    idx_3 = xml.find(">3</w:t>")
    idx_ins = xml.find('<w:t xml:space="preserve"> </w:t>')
    if idx_ins == -1: idx_ins = xml.find('<w:t> </w:t>')
        
    assert idx_art < idx_ins < idx_3, f"Order wrong! Art:{idx_art}, Ins:{idx_ins}, 3:{idx_3}"

def test_insertion_at_start_of_document():
    """
    Tests inserting text at the very beginning of the document.
    Original: "Contract"
    Modified: "Big Contract"
    """
    doc = Document()
    doc.add_paragraph("Contract")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    original_text = extract_text_from_stream(stream)
    modified_text = "Big " + original_text
    
    # We use generate_edits_from_text to test the Diff logic specifically
    from adeu.diff import generate_edits_from_text
    edits = generate_edits_from_text(original_text, modified_text)
    
    assert len(edits) > 0, "Should generate an edit for start-of-doc insertion"
    assert "Big" in edits[0].new_text
    assert edits[0].target_text in original_text

def test_insertion_multiple_splits_same_run():
    """
    Tests applying multiple insertions into the SAME original run.
    Original: "ARTICLE3 FEES"
    Edits: Insert " " after "ARTICLE", Insert " " after "3".
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ARTICLE3 FEES")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    # 1. Insert after ARTICLE
    e1 = DocumentEdit(operation=EditOperationType.INSERTION, target_text="ARTICLE", new_text=" ")
    # 2. Insert after 3 (Note: "3" is now in a split run if processed second? or first?)
    # Text context is tricky. Anchor is just "3".
    e2 = DocumentEdit(operation=EditOperationType.INSERTION, target_text="3", new_text=" ")
    
    engine = RedlineEngine(stream)
    engine.apply_edits([e1, e2])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    # Expected: ARTICLE <ins> </ins> 3 <ins> </ins> FEES
    # Just check that we have two insertions and text is split
    assert xml.count("<w:ins") == 2

def test_complex_run_sequence_repro():
    """
    Reproduction of 'ARTICLE3 FEESAN D PAYMENT' bug.
    Runs: ['ARTICLE3 FEES', 'AN', 'D', 'PAYMENT']
    Edits: Insert ' ' after 'ARTICLE3 FEES', Insert ' ' after 'AND'.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ARTICLE3 FEES")
    p.add_run("AN")
    p.add_run("D")
    p.add_run("PAYMENT")
    # NOTE: These runs have identical formatting (default).
    # They WILL be coalesced into "ARTICLE3 FEESANDPAYMENT" by RedlineEngine.
    # This tests the "Splitting a merged run" logic.
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    # Edit 1: Insert space after "ARTICLE3 FEES"
    e1 = DocumentEdit(operation=EditOperationType.INSERTION, target_text="ARTICLE3 FEES", new_text=" ")
    # Edit 2: Insert space after "AND" (which spans AN + D)
    e2 = DocumentEdit(operation=EditOperationType.INSERTION, target_text="AND", new_text=" ")
    
    engine = RedlineEngine(stream)
    engine.apply_edits([e1, e2])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    # Verify INS1 is between FEES and AN
    
    # Robust check using find
    idx_fees = xml.find("ARTICLE3 FEES")
    idx_ins1 = xml.find('w:id="1"')
    idx_an = xml.find(">AND</w:t>")
    
    assert idx_fees != -1
    assert idx_ins1 != -1
    assert idx_an != -1
    
    assert idx_fees < idx_ins1 < idx_an, f"Order Mismatch: FEES({idx_fees}) < INS1({idx_ins1}) < AN({idx_an})"

def test_overlapping_run_boundaries():
    """
    Test where target text ends exactly at run boundary, ensuring next run isn't grabbed.
    Run 1: "HELLO"
    Run 2: "WORLD"
    Target: "HELLO"
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("HELLO")
    p.add_run("WORLD")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    mapper = DocumentMapper(Document(stream))
    runs = mapper.find_target_runs("HELLO")
    
    assert len(runs) == 1
    assert runs[0].text == "HELLO"

def test_split_run_ordering_repro():
    """
    Reproduction of 'END0' bug.
    Original: 'e0'.
    Edit 1: Delete 'e'.
    Edit 2: Insert ' END' at end.
    """
    doc = Document()
    # Ensure clean slate
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
        
    p = doc.add_paragraph()
    p.add_run("e0")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    # Indices: e=0, 0=1. End=2.
    e1 = DocumentEdit(operation=EditOperationType.INSERTION, target_text="", new_text=" END")
    e1._match_start_index = 2
    e2 = DocumentEdit(operation=EditOperationType.DELETION, target_text="e", new_text=None)
    e2._match_start_index = 0
    
    engine = RedlineEngine(stream)
    # Engine sorts by index DESC, so e1 (2) applied first, then e2 (0).
    engine.apply_edits([e2, e1])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    # Expected: <del>e</del> <r>0</r> <ins> END</ins>
    idx_0 = xml.find(">0</w:t>")
    idx_ins = xml.find("> END</w:t>")
    
    assert idx_0 < idx_ins, f"0 ({idx_0}) should be before END ({idx_ins})"

def test_manual_context_disambiguation():
    """
    Proves that users can disambiguate targets simply by including more text,
    without needing specific 'context_before' fields.
    Scenario: "Section 1: Fee" ... "Section 2: Fee"
    Action: Change second Fee to Price by targeting "Section 2: Fee".
    """
    doc = Document()
    doc.add_paragraph("Section 1: Fee")
    doc.add_paragraph("Section 2: Fee")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    # 1. Ambiguous Edit (would normally hit the first one or warn)
    # But here we use the "Context Strategy"
    edit = DocumentEdit(
        operation=EditOperationType.MODIFICATION,
        target_text="Section 2: Fee",
        new_text="Section 2: Price",
        comment="Disambiguated via context"
    )
    
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    
    assert applied == 1
    assert skipped == 0
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    # Verify Section 1 is untouched
    assert "Section 1: Fee" in xml or "Section 1: </w:t><w:t>Fee" in xml
    
    # Verify Section 2 is Redlined
    # We expect "Section 2: Fee" to be deleted and "Section 2: Price" inserted
    assert "<w:delText>Section 2: Fee</w:delText>" in xml or ("Section 2: " in xml and "Fee</w:delText>" in xml)
    assert "<w:t>Section 2: Price</w:t>" in xml